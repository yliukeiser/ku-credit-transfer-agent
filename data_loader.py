"""
Loads and indexes KU program requirements from Excel and transfer policy from Word.
"""
import json
import re
from pathlib import Path
from typing import List, Optional

import openpyxl
from docx import Document

# ── paths ──────────────────────────────────────────────────────────────────────
# Look for data files in the project's /data folder first,
# then fall back to ~/Downloads (for local development)
_HERE      = Path(__file__).parent
_DATA_DIR  = _HERE / "data"
_DOWNLOADS = Path.home() / "Downloads"

def _find(filename: str) -> Path:
    if (_DATA_DIR / filename).exists():
        return _DATA_DIR / filename
    return _DOWNLOADS / filename

REQUIREMENTS_XLSX  = _find("ProgramCourseReqIndexTable.xlsx")
POLICY_DOCX        = _find("Undergraduate Transfer of Credit Policy.docx")
SPECIAL_RULES_XLSX = _find("ProgramTransferPolicySpecialRules.xlsx")


# ── program requirements ───────────────────────────────────────────────────────

def load_program_requirements() -> dict:
    """
    Returns a nested dict:
      {
        "Accounting|Associate of Arts": {
          "program_name": "Accounting",
          "degree_type": "Associate of Arts",
          "degree_type_short": "AA",
          "full_name": "Associate of Arts in Accounting",
          "program_code": "KC-AAAC",
          "categories": {
            "Accounting Major Courses": {
              "total_credits": 24,
              "disciplines": {
                "Accounting Major Courses": {
                  "total_credits": 24,
                  "courses": [
                    {"code": "ACG1001", "name": "Accounting Principles I", "credits": 3},
                    ...
                  ]
                }
              }
            }
          }
        }
      }
    """
    wb = openpyxl.load_workbook(REQUIREMENTS_XLSX)
    ws = wb["Sheet1"]

    programs: dict = {}

    for row in ws.iter_rows(min_row=2, values_only=True):
        (prog_name, deg_type, deg_short, full_name, prog_ver_id,
         prog_code, category, cat_credits, discipline, disc_credits,
         course_code, course_name, course_credits) = row

        if not prog_name:
            continue

        key = f"{prog_name}|{deg_type}"
        if key not in programs:
            programs[key] = {
                "program_name": prog_name,
                "degree_type": deg_type,
                "degree_type_short": deg_short,
                "full_name": full_name,
                "program_code": prog_code,
                "categories": {}
            }

        cats = programs[key]["categories"]
        if category not in cats:
            cats[category] = {
                "total_credits": cat_credits,
                "disciplines": {}
            }

        discs = cats[category]["disciplines"]
        if discipline not in discs:
            discs[discipline] = {
                "total_credits": disc_credits,
                "courses": []
            }

        if course_code:
            try:
                credits_val = int(course_credits) if course_credits else 0
            except (ValueError, TypeError):
                credits_val = 0
            discs[discipline]["courses"].append({
                "code": str(course_code).strip(),
                "name": course_name,
                "credits": credits_val,
            })

    return programs


def list_programs(programs: dict) -> List[dict]:
    """Returns a flat sorted list of {key, program_name, degree_type, full_name}."""
    result = []
    for key, info in programs.items():
        result.append({
            "key": key,
            "program_name": info["program_name"],
            "degree_type": info["degree_type"],
            "full_name": info["full_name"],
            "program_code": info["program_code"],
        })
    return sorted(result, key=lambda x: (x["degree_type"], x["program_name"]))


def search_programs(programs: dict, query: str) -> List[dict]:
    """Case-insensitive substring search across program name and full name."""
    q = query.lower()
    return [
        p for p in list_programs(programs)
        if q in p["program_name"].lower()
        or q in p["full_name"].lower()
        or q in p["degree_type"].lower()
    ]


def get_program(programs: dict, key: str) -> Optional[dict]:
    return programs.get(key)


# ── transfer credit policy ─────────────────────────────────────────────────────

def load_transfer_policy() -> dict:
    """
    Returns a dict with:
      - 'full_text': complete policy as a single string
      - 'rules': list of key policy rule strings
      - 'aice_equivalencies': list of {exam, ku_courses, credits}
    """
    doc = Document(POLICY_DOCX)

    # Build full text
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # Extract AICE equivalency table
    aice_equiv = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if "AICE" in headers[0]:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3 and cells[0]:
                    aice_equiv.append({
                        "exam": cells[0],
                        "ku_courses": cells[1],
                        "credits": cells[2]
                    })

    # Key distilled rules (verbatim or paraphrased from policy)
    rules = [
        "Minimum grade of C (2.0 on a 4.0 scale) required for any course to be considered for transfer.",
        "Transfer credits are granted only for courses applicable to the student's program of study.",
        "Credits from regionally or nationally accredited institutions are accepted; non-accredited sources require additional review.",
        "Students must complete the final 25% of their program at Keiser University.",
        "Clock-hour conversions: 15 lecture hours = 1 credit; 30 lab hours = 1 credit; 45 externship hours = 1 credit.",
        "International transcripts are evaluated on a course-equivalency basis by an approved evaluation service.",
        "Students with an Associate of Arts from a Florida CCNS-aligned institution have lower-division General Education requirements waived.",
        "Students with an AA from a Florida public community college (2.0+ GPA) under the Statewide Articulation Agreement have General Education requirements waived.",
        "Students with any Bachelor's degree from a USDE-recognized institution have all General Education requirements waived.",
        "Approved AICE examination scores (A, B, C, D, or E on A and AS levels) earn Keiser University course credit as listed in the AICE equivalency table.",
        "Course descriptions from prior institutions are analyzed to confirm content alignment before credit is accepted.",
        "Individual programmatic requirements supersede general education transfer guidelines.",
        "Articulation agreements between Keiser University and other colleges are recognized for transfer of credit.",
        "Veterans must report all previous education and training; Keiser evaluates military transcripts (JST) per VA guidelines.",
        "Prior Learning Assessment (PLA) credit may be awarded based on examination results from CLEP, DANTES, AP, or other approved programs.",
    ]

    return {
        "full_text": full_text,
        "rules": rules,
        "aice_equivalencies": aice_equiv,
    }


# ── special rules index ────────────────────────────────────────────────────────

def load_special_rules_index() -> List[dict]:
    """
    Returns list of {degree_type, program_name, start_page, end_page, chunk}
    for advisor reference (the special rules Excel is a page index into a
    separate policy PDF; we store it so the agent can cite the source).
    """
    wb = openpyxl.load_workbook(SPECIAL_RULES_XLSX)
    ws = wb["Sheet1"]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        deg_type, prog_name, start, end, chunk = row
        if prog_name:
            rows.append({
                "degree_type": deg_type,
                "program_name": prog_name,
                "start_page": start,
                "end_page": end,
                "chunk": chunk,
            })
    return rows


def find_special_rules(index: List[dict], program_name: str) -> Optional[dict]:
    """Find special-rules page reference for a given program name."""
    q = program_name.lower()
    for entry in index:
        if entry["program_name"].lower() == q:
            return entry
    return None


# ── singleton cache ────────────────────────────────────────────────────────────

_cache: dict = {}

def get_data() -> dict:
    """Load all data once and cache in memory."""
    if not _cache:
        _cache["programs"] = load_program_requirements()
        _cache["policy"]   = load_transfer_policy()
        _cache["special_rules_index"] = load_special_rules_index()
    return _cache
